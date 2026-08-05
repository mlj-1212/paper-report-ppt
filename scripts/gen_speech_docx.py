#!/usr/bin/env python3
"""
Speech script DOCX generator for paper-report-ppt v4.0.

Reads a structured JSON file (speech_data.json) produced by the AI agent
and renders a polished DOCX speech script with:
  - Heading hierarchy (H1/H2/H3)
  - CJK font (Microsoft YaHei) via XML eastAsia attribute
  - Body text 12pt, 1.5x line spacing
  - Duration allocation table with light green header
  - Centered page-number footer
  - A4 page size, 1-inch margins

Usage:
  python gen_speech_docx.py --input speech_data.json
  # --output 可省略：省略时按标题自动命名为 <标题>_演讲稿.docx（落在 speech_data.json 同目录）
  python gen_speech_docx.py -i speech_data.json -o speech.docx --verbose

Input JSON schema (speech_data.json):
{
  "title": "文献汇报演讲稿：<论文标题>",
  "meta": {
    "literature": "作者, 年份, 期刊",
    "scenario": "研究生组会",
    "duration_minutes": 20,
    "date": "2026-07-29"
  },
  "opening": "开场白全文...",
  "sections": [
    {
      "part_title": "第一部分：研究背景",
      "pages": [
        {
          "page_num": "P03",
          "page_title": "研究背景：RSV与植物防御",
          "duration_minutes": 1.5,
          "content": "该页的完整口头讲解文字..."
        }
      ]
    }
  ],
  "closing": "结束语全文...",
  "duration_table": [
    {"part": "开场白", "pages": "-", "duration": "1分钟"},
    {"part": "第一部分", "pages": "P03-P04", "duration": "3分钟"}
  ],
  "tips": ["时长分配建议", "重点强调提示", "可能被提问的预判"]
}
"""

import argparse
import json
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ═══════════════════════════════════════════════════════════
#  Constants
# ═══════════════════════════════════════════════════════════

CJK_FONT = "Microsoft YaHei"
LATIN_FONT = "Arial"
BODY_SIZE_PT = 12
BODY_LINE_SPACING = 1.5
TABLE_HEADER_BG = "EDF5E8"  # light green
TABLE_HEADER_TEXT = RGBColor(0x2C, 0x3E, 0x50)


# ═══════════════════════════════════════════════════════════
#  XML helpers
# ═══════════════════════════════════════════════════════════

def _set_run_cjk_font(run, font_name=CJK_FONT, latin_font=LATIN_FONT):
    """Set CJK eastAsia font on a run via XML, ensuring Chinese renders correctly."""
    rPr = run._r.get_or_add_rPr()
    # Remove existing eastAsia elements
    for ea in rPr.findall(qn("w:rFonts")):
        rPr.remove(ea)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), latin_font)
    rFonts.set(qn("w:hAnsi"), latin_font)
    rFonts.set(qn("w:eastAsia"), font_name)
    rPr.insert(0, rFonts)


def _set_cell_shading(cell, fill_hex):
    """Set table cell background color via XML."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _add_page_number_field(paragraph):
    """Add a centered '第 N 页' page number field to a footer paragraph."""
    run = paragraph.add_run("第 ")
    _set_run_cjk_font(run)
    run.font.size = Pt(10)

    # PAGE field
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run2 = paragraph.add_run()
    _set_run_cjk_font(run2)
    run2.font.size = Pt(10)
    run2._r.append(fld_begin)
    run2._r.append(instr)
    run2._r.append(fld_end)

    run3 = paragraph.add_run(" 页")
    _set_run_cjk_font(run3)
    run3.font.size = Pt(10)


# ═══════════════════════════════════════════════════════════
#  Document setup
# ═══════════════════════════════════════════════════════════

def _setup_page(doc):
    """Configure A4 page size and 1-inch margins."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4 width
    section.page_height = Cm(29.7)  # A4 height
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Footer with centered page number
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(footer_para)


def _add_heading(doc, text, level):
    """Add a heading with CJK font support."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        _set_run_cjk_font(run)
    return heading


def _add_body_paragraph(doc, text):
    """Add a body paragraph with 12pt, 1.5x spacing, CJK font."""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = BODY_LINE_SPACING
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.size = Pt(BODY_SIZE_PT)
    _set_run_cjk_font(run)
    return para


def _add_meta_line(doc, label, value):
    """Add a metadata line (italic, smaller)."""
    para = doc.add_paragraph()
    run_label = para.add_run(f"{label}：")
    run_label.font.size = Pt(11)
    run_label.bold = True
    _set_run_cjk_font(run_label)
    run_value = para.add_run(value)
    run_value.font.size = Pt(11)
    run_value.italic = True
    _set_run_cjk_font(run_value)
    return para


def _add_duration_table(doc, duration_rows):
    """Add a 3-column duration allocation table with light green header."""
    if not duration_rows:
        return

    table = doc.add_table(rows=1 + len(duration_rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    headers = ["部分", "页面", "预计时长"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = TABLE_HEADER_TEXT
        _set_run_cjk_font(run)
        _set_cell_shading(cell, TABLE_HEADER_BG)

    # Data rows
    for row_idx, row_data in enumerate(duration_rows):
        cells = table.rows[row_idx + 1].cells
        values = [
            row_data.get("part", ""),
            row_data.get("pages", ""),
            row_data.get("duration", ""),
        ]
        for i, val in enumerate(values):
            cells[i].text = ""
            para = cells[i].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(val))
            run.font.size = Pt(11)
            _set_run_cjk_font(run)

    # Add spacing after table
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════
#  Main generation logic
# ═══════════════════════════════════════════════════════════

def generate_speech_docx(data, output_path):
    """Generate a DOCX speech script from structured JSON data."""
    doc = Document()

    # Page setup + footer
    _setup_page(doc)

    # --- Title (H1) ---
    title = data.get("title", "文献汇报演讲稿")
    _add_heading(doc, title, level=1)

    # --- Metadata block ---
    meta = data.get("meta", {})
    if meta:
        _add_meta_line(doc, "文献", meta.get("literature", "N/A"))
        _add_meta_line(doc, "汇报场景", meta.get("scenario", "研究生组会"))
        duration = meta.get("duration_minutes", 20)
        _add_meta_line(doc, "预计时长", f"约 {duration} 分钟")
        report_date = meta.get("date", date.today().isoformat())
        _add_meta_line(doc, "生成日期", report_date)

    # Horizontal rule (using a paragraph border)
    hr_para = doc.add_paragraph()
    pPr = hr_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # --- Opening ---
    opening = data.get("opening", "")
    if opening:
        _add_heading(doc, "开场白", level=2)
        _add_body_paragraph(doc, opening)

    # --- Sections ---
    sections = data.get("sections", [])
    for section in sections:
        part_title = section.get("part_title", "")
        if part_title:
            _add_heading(doc, part_title, level=2)

        pages = section.get("pages", [])
        for page in pages:
            page_title = page.get("page_title", "")
            page_num = page.get("page_num", "")
            header_text = f"{page_num} — {page_title}" if page_num else page_title
            if header_text:
                _add_heading(doc, header_text, level=3)

            content = page.get("content", "")
            if content:
                _add_body_paragraph(doc, content)

            duration = page.get("duration_minutes")
            if duration:
                _add_meta_line(doc, "预计时长", f"{duration} 分钟")

    # --- Closing ---
    closing = data.get("closing", "")
    if closing:
        _add_heading(doc, "结束语", level=2)
        _add_body_paragraph(doc, closing)

    # --- Duration table ---
    duration_table = data.get("duration_table", [])
    if duration_table:
        _add_heading(doc, "演讲时长分配", level=2)
        _add_duration_table(doc, duration_table)

    # --- Tips ---
    tips = data.get("tips", [])
    if tips:
        _add_heading(doc, "演讲提示", level=2)
        for tip in tips:
            _add_body_paragraph(doc, f"• {tip}")

    # Save
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ═══════════════════════════════════════════════════════════
#  CLI
# ═══════════════════════════════════════════════════════════

# ── 输出文件名自动命名（按文献命名，用户好找）──────────────────────────
def _safe_stem(name, max_len=40):
    """把任意标题/文件名转成安全的文件 stem（保留中英文、数字、下划线、连字符）。"""
    if not name:
        return ""
    name = re.sub(r'^(文献汇报演讲稿[:：]?|文献汇报[:：]?|文献精读汇报[:：]?)', '', name.strip())
    name = re.sub(r'\.(pdf|pptx|docx)$', '', name, flags=re.I)
    keep = []
    for ch in name:
        if ch.isalnum() or '\u4e00' <= ch <= '\u9fff' or ch in ' _-':
            keep.append(ch)
        else:
            keep.append(' ')
    s = ''.join(keep)
    s = re.sub(r'\s+', '_', s.strip()).strip('_-')
    if len(s) > max_len:
        s = s[:max_len].rstrip('_-')
    return s


def main():
    parser = argparse.ArgumentParser(
        description="Generate a DOCX speech script from structured JSON data."
    )
    parser.add_argument(
        "--input", "-i", required=True,
        help="Path to speech_data.json"
    )
    parser.add_argument(
        "--output", "-o", default=None,
        help="输出 DOCX 路径；缺省时按文献标题自动命名（<标题>_演讲稿.docx，落在输入 json 同目录）"
    )
    parser.add_argument(
        "--stem", default=None,
        help="文献简称，用于自动命名输出文件；缺省时从标题推导"
    )
    parser.add_argument(
        "--verbose", "-v", action="store_true",
        help="Print detailed progress"
    )
    args = parser.parse_args()

    # Load JSON
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"[ERROR] Input file not found: {input_path}", file=sys.stderr)
        return 1

    with open(input_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    if args.verbose:
        print(f"[INFO] Loaded speech data from {input_path}")
        print(f"[INFO] Title: {data.get('title', 'N/A')}")
        print(f"[INFO] Sections: {len(data.get('sections', []))}")

    # ── 决定输出文件名（缺省时按文献标题自动命名，落在输入 json 同目录）──
    if args.output:
        output_path = Path(args.output)
    else:
        _stem = args.stem or _safe_stem(data.get("title", "")) or "文献汇报演讲稿"
        output_path = input_path.parent / f"{_stem}_演讲稿.docx"

    # Generate DOCX
    try:
        output_path = generate_speech_docx(data, output_path)
        if args.verbose:
            print(f"[INFO] DOCX saved to {output_path}")
        print(f"✅ 演讲稿已生成: {output_path}")
        return 0
    except Exception as e:
        print(f"[ERROR] Failed to generate DOCX: {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
